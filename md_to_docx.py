import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(md_file, docx_file):
    # Criar um novo documento Word
    doc = docx.Document()
    
    # Configurar margens do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Ler o conteúdo do arquivo markdown
    with open(md_file, 'r', encoding='utf-8') as file:
        md_content = file.readlines()
    
    # Processar o conteúdo linha por linha
    current_list = []
    in_code_block = False
    
    for line in md_content:
        line = line.rstrip()
        
        # Verificar se é um bloco de código
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        # Pular processamento dentro de blocos de código
        if in_code_block:
            p = doc.add_paragraph()
            p.add_run(line).font.name = 'Courier New'
            p.add_run().add_break()
            continue
        
        # Título principal (# )
        if line.startswith('# '):
            title = doc.add_heading(line[2:], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Título nível 2 (## )
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        
        # Título nível 3 (### )
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        
        # Título nível 4 (#### )
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 3)
        
        # Item de lista (- ou * )
        elif line.startswith('- ') or line.startswith('* '):
            current_list.append(line[2:])
        
        # Linha em branco - finaliza lista se houver
        elif line.strip() == '':
            if current_list:
                for item in current_list:
                    p = doc.add_paragraph(item)
                    p.style = 'List Bullet'
                current_list = []
            else:
                doc.add_paragraph()
        
        # Texto normal
        else:
            if current_list:
                for item in current_list:
                    p = doc.add_paragraph(item)
                    p.style = 'List Bullet'
                current_list = []
            
            doc.add_paragraph(line)
    
    # Salvar o documento
    doc.save(docx_file)
    print(f"Documento Word criado com sucesso: {docx_file}")

if __name__ == "__main__":
    create_word_document('/home/ubuntu/projeto_sqlite/documentacao.md', 
                         '/home/ubuntu/projeto_sqlite/documentacao.docx')
